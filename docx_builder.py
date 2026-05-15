"""
docx_builder.py
NUS Development Office – Prospect Profile Generator
Builds the .docx profile files using python-docx + direct XML manipulation.
"""

import io
from typing import Optional, List, Dict

from docx import Document
from docx.shared import Pt, Inches, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image


# ──────────────────────────────────────────────
# Constants (all widths in DXA / twips, 1 inch = 1440 DXA)
# ──────────────────────────────────────────────
LABEL_W   = 2380   # label column
VALUE_W   = 6526   # value column
TABLE_W   = 8906   # total table width
TABLE_IND = 120    # left indent of table from body text
FONT_NAME = "Roboto"
FONT_PT   = 11


# ══════════════════════════════════════════════
# DocxBuilder class
# ══════════════════════════════════════════════

class DocxBuilder:

    def __init__(self):
        self.doc = Document()
        self._configure_page()
        self._setup_header()
        self._setup_footer()

    # ──────────────────────────────────────────
    # Page / section setup
    # ──────────────────────────────────────────

    def _configure_page(self):
        """A4 page, 1-inch margins all round."""
        sec = self.doc.sections[0]
        sec.page_width   = Twips(11906)
        sec.page_height  = Twips(16838)
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1)
        sec.right_margin  = Inches(1)

    def _setup_header(self):
        """Bold 'DEVELOPMENT OFFICE' left, 'CONFIDENTIAL' right."""
        header = self.doc.sections[0].header
        p = header.paragraphs[0]
        p.clear()

        # Tab stop flush-right at table width
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab  = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:pos"), str(TABLE_W + TABLE_IND))
        tabs.append(tab)
        pPr.append(tabs)

        r1 = p.add_run("DEVELOPMENT OFFICE")
        r1.bold = True
        r1.font.name = FONT_NAME
        r1.font.size = Pt(FONT_PT)

        r2 = p.add_run("\tCONFIDENTIAL")
        r2.font.name = FONT_NAME
        r2.font.size = Pt(FONT_PT)

        # Thin bottom rule under header
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "4")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "000000")
        pBdr.append(bot)
        pPr.append(pBdr)

    def _setup_footer(self):
        """Centred page number '1'."""
        footer = self.doc.sections[0].footer
        p = footer.paragraphs[0]
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("1")
        r.font.name = FONT_NAME
        r.font.size = Pt(FONT_PT)

    # ──────────────────────────────────────────
    # Confidentiality block
    # ──────────────────────────────────────────

    def _add_conf_lines(self):
        for text, bold, italic in [
            ("NUS CONFIDENTIAL",           True,  False),
            ("for NUS internal use only",  False, True),
            ("NOT for external circulation", False, True),
        ]:
            p = self.doc.add_paragraph()
            _para_spacing(p, before=0, after=0)
            r = p.add_run(text)
            r.bold   = bold
            r.italic = italic
            r.font.name = FONT_NAME
            r.font.size = Pt(FONT_PT)

    # ──────────────────────────────────────────
    # Photo block
    # ──────────────────────────────────────────

    def _add_photo(self, photo_bytes: Optional[bytes], label: str = "PHOTO"):
        p = self.doc.add_paragraph()
        _para_spacing(p, before=6, after=6)

        if photo_bytes:
            try:
                img = Image.open(io.BytesIO(photo_bytes))
                # Normalise orientation
                img_bytes = io.BytesIO()
                img.save(img_bytes, format="PNG")
                img_bytes.seek(0)
                p.add_run().add_picture(img_bytes, height=Inches(1.8))
                return
            except Exception:
                pass

        # Placeholder
        r = p.add_run(f"[{label}]")
        r.italic = True
        r.font.name = FONT_NAME
        r.font.size = Pt(FONT_PT)

    # ──────────────────────────────────────────
    # Table helpers
    # ──────────────────────────────────────────

    def _make_table(self, n_rows: int):
        """Return a styled 2-column table."""
        tbl = self.doc.add_table(rows=n_rows, cols=2)
        _set_table_props(tbl)
        return tbl

    def _fill_row(self, row, label: str, value_fn):
        """Populate one table row: label in col-0, content built by value_fn in col-1."""
        lc = row.cells[0]
        vc = row.cells[1]
        _set_cell_props(lc, LABEL_W)
        _set_cell_props(vc, VALUE_W)
        # Label – plain text, no bullet
        _clear_para(lc.paragraphs[0])
        _plain_run(lc.paragraphs[0], label)
        # Value
        value_fn(vc)

    # ──────────────────────────────────────────
    # Public build methods
    # ──────────────────────────────────────────

    def build_individual(self, data: dict, photo_bytes: Optional[bytes] = None) -> bytes:
        """Build an individual prospect profile and return .docx bytes."""
        pronoun = "She" if data.get("gender", "male").lower() == "female" else "He"

        self._add_conf_lines()
        self._add_photo(photo_bytes, label="PHOTO")

        include_adverse = bool(data.get("adverse_news"))
        n_rows = 11 + (1 if include_adverse else 0)
        tbl = self._make_table(n_rows)
        rows = tbl.rows

        # Row 0 – Name
        self._fill_row(rows[0], "Name:", lambda c: _bullets(c, [data.get("name", "")]))

        # Row 1 – Key position
        self._fill_row(rows[1], "Key position/organisation:",
                       lambda c: _bullets(c, [data.get("key_position", "Not publicly available.")]))

        # Row 2 – Age
        self._fill_row(rows[2], "Age:",
                       lambda c: _bullets(c, [data.get("age", "Not publicly available.")]))

        # Row 3 – Nationality
        self._fill_row(rows[3], "Nationality:",
                       lambda c: _bullets(c, [data.get("nationality", "Not publicly available.")]))

        # Row 4 – Net Worth
        self._fill_row(rows[4], "Net Worth:",
                       lambda c: _bullets(c, [data.get("net_worth", "Not publicly available.")]))

        # Row 5 – Education
        self._fill_row(rows[5], "Education:",
                       lambda c: _bullets(c, data.get("education") or ["Not publicly available."]))

        # Row 6 – Brief Biography
        def bio_fn(c):
            intro = data.get("biography_intro") or []
            cur   = data.get("biography_current_positions") or []
            past  = data.get("biography_past_positions") or []
            fam   = data.get("biography_family") or []

            all_items = []
            all_items += [(b, False) for b in intro]
            if cur:
                all_items.append((f"{pronoun}s other current positions include:", True))
                all_items += [(b, False) for b in cur]
            if past:
                all_items.append((f"{pronoun}s notable past positions include:", True))
                all_items += [(b, False) for b in past]
            if fam:
                all_items.append(("Family:", True))
                all_items += [(b, False) for b in fam]

            if not all_items:
                all_items = [("Not publicly available.", False)]

            _bullets_mixed(c, all_items)

        self._fill_row(rows[6], "Brief biography:", bio_fn)

        # Row 7 – Giving
        self._fill_row(rows[7], "Giving to other organisations:",
                       lambda c: _bullets(c, data.get("giving") or ["Not publicly available."]))

        # Row 8 – Demonstrated interests
        self._fill_row(rows[8], "Demonstrated interests:",
                       lambda c: _bullets(c, data.get("interests") or ["Not publicly available."]))

        # Row 9 – Other interesting facts
        def facts_fn(c):
            awards = data.get("awards") or []
            other  = data.get("other_facts") or []
            items  = []
            if awards:
                items.append(("Awards:", True))
                items += [(a, False) for a in awards]
            if other:
                if awards:
                    items.append(("Other:", True))
                items += [(o, False) for o in other]
            if not items:
                items = [("Not publicly available.", False)]
            _bullets_mixed(c, items)

        self._fill_row(rows[9], "Other interesting facts:", facts_fn)

        # Row 10 – Adverse news (conditional)
        row_idx = 10
        if include_adverse:
            adverse_items = data.get("adverse_news", []) + [
                "Note: This is a preliminary search only and may not reflect the complete list of adverse news."
            ]
            self._fill_row(rows[row_idx], "Adverse news:",
                           lambda c: _bullets(c, adverse_items))
            row_idx += 1

        # Second table – connector + gift ideas
        tbl2 = self._make_table(2)

        # Connector row
        def connector_fn(c):
            connectors = data.get("connectors") or []
            if not connectors:
                _bullets(c, ["No suitable connector identified by Claude from public sources."])
                return
            items: List[tuple] = []
            for conn in connectors[:5]:
                if not isinstance(conn, dict):
                    continue
                name_title = conn.get("name_title", "")
                rel        = conn.get("relationship_to_prospect", "")
                nus        = conn.get("nus_connection", "")
                approach   = conn.get("recommended_approach", "")
                items.append((name_title, True))
                if rel:
                    items.append((f"Relationship to prospect: {rel}", False))
                if nus:
                    items.append((f"NUS connection: {nus}", False))
                if approach:
                    items.append((f"Recommended approach: {approach}", False))
            _bullets_mixed(c, items)

        self._fill_row(tbl2.rows[0],
                       "Potential Connector to connect with NUS and Contact Details:",
                       connector_fn)

        # Gift ideas row
        self._fill_row(tbl2.rows[1], "Gift Ideas for Donations to NUS:",
                       lambda c: _bullets(c, data.get("gift_ideas") or ["Not publicly available."]))

        # Sources footnote
        sources = data.get("sources") or []
        if sources:
            self._add_sources(sources)

        return self._to_bytes()

    def build_company(self, data: dict, photo_bytes: Optional[bytes] = None) -> bytes:
        """Build a company prospect profile and return .docx bytes."""
        org_name = data.get("organisation_name", "the organisation")

        self._add_conf_lines()
        self._add_photo(photo_bytes, label="LOGO")

        include_adverse = bool(data.get("adverse_news"))
        n_rows = 9 + (1 if include_adverse else 0)
        tbl = self._make_table(n_rows)
        rows = tbl.rows

        # Row 0 – Organisation Name
        self._fill_row(rows[0], "Organisation Name:",
                       lambda c: _bullets(c, [data.get("organisation_name", "")]))

        # Row 1 – Year of Establishment
        self._fill_row(rows[1], "Year of Establishment:",
                       lambda c: _bullets(c, [data.get("year_established", "Not publicly available.")]))

        # Row 2 – Country of Registration
        self._fill_row(rows[2], "Country of Registration:",
                       lambda c: _bullets(c, [data.get("country_of_registration", "Not publicly available.")]))

        # Row 3 – Annual Revenue
        self._fill_row(rows[3], "Annual Revenue:",
                       lambda c: _bullets(c, [data.get("annual_revenue", "Not publicly available.")]))

        # Row 4 – Brief biography
        def bio_fn(c):
            intro      = data.get("biography_intro") or []
            subsections = data.get("biography_subsections") or []
            items: List[tuple] = [(b, False) for b in intro]
            for sub in subsections:
                if not isinstance(sub, dict):
                    continue
                label   = sub.get("label", "")
                bullets = sub.get("bullets") or []
                if label:
                    items.append((label, True))
                items += [(b, False) for b in bullets]
            if not items:
                items = [("Not publicly available.", False)]
            _bullets_mixed(c, items)

        self._fill_row(rows[4], "Brief biography:", bio_fn)

        # Row 5 – Giving
        self._fill_row(rows[5], "Giving to other organisations:",
                       lambda c: _bullets(c, data.get("giving") or ["Not publicly available."]))

        # Row 6 – Demonstrated interests
        self._fill_row(rows[6], "Demonstrated interests:",
                       lambda c: _bullets(c, data.get("interests") or ["Not publicly available."]))

        # Row 7 – Other interesting facts
        def facts_fn(c):
            awards = data.get("awards") or []
            other  = data.get("other_facts") or []
            items: List[tuple] = []
            if awards:
                items.append(("Awards:", True))
                items += [(a, False) for a in awards]
            if other:
                if awards:
                    items.append(("Other:", True))
                items += [(o, False) for o in other]
            if not items:
                items = [("Not publicly available.", False)]
            _bullets_mixed(c, items)

        self._fill_row(rows[7], "Other interesting facts:", facts_fn)

        # Row 8 – Adverse news (conditional)
        row_idx = 8
        if include_adverse:
            adverse_items = data.get("adverse_news", []) + [
                "Note: This is a preliminary search only and may not reflect the complete list of adverse news."
            ]
            self._fill_row(rows[row_idx], "Adverse news:",
                           lambda c: _bullets(c, adverse_items))
            row_idx += 1

        # Second table
        tbl2 = self._make_table(2)

        def connector_fn(c):
            connectors = data.get("connectors") or []
            if not connectors:
                _bullets(c, ["No suitable connector identified by Claude from public sources."])
                return
            items: List[tuple] = []
            for conn in connectors[:5]:
                if not isinstance(conn, dict):
                    continue
                name_title = conn.get("name_title", "")
                rel        = conn.get("relationship_to_prospect", "")
                nus        = conn.get("nus_connection", "")
                approach   = conn.get("recommended_approach", "")
                items.append((name_title, True))
                if rel:
                    items.append((f"Relationship to prospect: {rel}", False))
                if nus:
                    items.append((f"NUS connection: {nus}", False))
                if approach:
                    items.append((f"Recommended approach: {approach}", False))
            _bullets_mixed(c, items)

        self._fill_row(tbl2.rows[0],
                       "Potential Connector to connect with NUS and Contact Details:",
                       connector_fn)

        self._fill_row(tbl2.rows[1], "Gift Ideas for Donations to NUS:",
                       lambda c: _bullets(c, data.get("gift_ideas") or ["Not publicly available."]))

        sources = data.get("sources") or []
        if sources:
            self._add_sources(sources)

        return self._to_bytes()

    # ──────────────────────────────────────────
    # Sources footnote
    # ──────────────────────────────────────────

    def _add_sources(self, sources: List[str]):
        p0 = self.doc.add_paragraph()
        _para_spacing(p0, before=8, after=0)
        r0 = p0.add_run("Sources:")
        r0.bold = True
        r0.italic = True
        r0.font.name = FONT_NAME
        r0.font.size = Pt(9)

        for i, src in enumerate(sources, 1):
            ps = self.doc.add_paragraph()
            _para_spacing(ps, before=0, after=0)
            rs = ps.add_run(f"{i}. {src}")
            rs.font.name = FONT_NAME
            rs.font.size = Pt(9)
            rs.italic = True

    # ──────────────────────────────────────────
    # Serialise
    # ──────────────────────────────────────────

    def _to_bytes(self) -> bytes:
        buf = io.BytesIO()
        self.doc.save(buf)
        buf.seek(0)
        return buf.getvalue()


# ══════════════════════════════════════════════
# Module-level XML helpers
# ══════════════════════════════════════════════

def _para_spacing(para, before: int = 0, after: int = 2):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)


def _plain_run(para, text: str, bold=False, italic=False, size=FONT_PT):
    """Add a plain (no-bullet) run to para."""
    r = para.add_run(text)
    r.bold      = bold
    r.italic    = italic
    r.font.name = FONT_NAME
    r.font.size = Pt(size)
    return r


def _clear_para(para):
    """Remove all runs from a paragraph without deleting the paragraph element."""
    p_el = para._p
    for child in list(p_el):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("r", "hyperlink", "ins", "del"):
            p_el.remove(child)


def _set_table_props(tbl_obj):
    """Apply width, indent, and strip jc from a python-docx Table."""
    tbl = tbl_obj._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # Width
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"),    str(TABLE_W))
    tblW.set(qn("w:type"), "dxa")

    # Indent
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"),    str(TABLE_IND))
    tblInd.set(qn("w:type"), "dxa")

    # Remove jc (no explicit alignment)
    jc = tblPr.find(qn("w:jc"))
    if jc is not None:
        tblPr.remove(jc)

    # Remove table-level borders so only cell-level borders are active
    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is not None:
        tblPr.remove(tblBorders)

    # Remove table look
    tblLook = tblPr.find(qn("w:tblLook"))
    if tblLook is not None:
        tblPr.remove(tblLook)


def _set_cell_props(cell, width_dxa: int):
    """Set cell width, borders (4 sides only, no inside rules), and padding."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Width
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"),    str(width_dxa))
    tcW.set(qn("w:type"), "dxa")

    # Borders
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is not None:
        tcPr.remove(tcBorders)
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        tcBorders.append(b)
    tcPr.append(tcBorders)

    # Internal padding
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is not None:
        tcPr.remove(tcMar)
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", 80), ("bottom", 80), ("left", 120), ("right", 120)):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"),    str(val))
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)

    # Remove shading
    shd = tcPr.find(qn("w:shd"))
    if shd is not None:
        tcPr.remove(shd)


def _bullet_para(para, text: str, italic=False, bold=False, sub_label=False):
    """Format an existing paragraph as a bullet or sub-label."""
    _clear_para(para)
    _para_spacing(para, before=0, after=1)

    if sub_label:
        # Italic sub-labels – no list numbering, just indent
        para.paragraph_format.left_indent  = Pt(0)
        para.paragraph_format.first_line_indent = Pt(0)
        r = para.add_run(text)
        r.italic    = True
        r.bold      = False
        r.font.name = FONT_NAME
        r.font.size = Pt(FONT_PT)
    else:
        # Use built-in List Bullet style
        try:
            from docx.opc.exceptions import PackageNotFoundError
            para.style = para._p.getroottree().getroot()  # will fail – handled below
        except Exception:
            pass
        # Set numbering via XML (most reliable cross-platform approach)
        _apply_list_bullet(para)
        r = para.add_run(text)
        r.italic    = italic
        r.bold      = bold
        r.font.name = FONT_NAME
        r.font.size = Pt(FONT_PT)


def _apply_list_bullet(para):
    """
    Apply a bullet list format to para via XML, using the document's
    built-in numbering or creating a minimal one if absent.
    """
    doc_part = para._p.getroottree().getroot()
    # Locate the Document element's parent part to access numbering
    # We will use the paragraph style approach: set pStyle to 'ListBullet'
    # and rely on the template's built-in numbering definition.
    pPr = para._p.get_or_add_pPr()

    # Set paragraph style to List Bullet via XML (avoids KeyError on some envs)
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        pStyle = OxmlElement("w:pStyle")
        pPr.insert(0, pStyle)
    pStyle.set(qn("w:val"), "ListBullet")

    # Explicit indent so it looks like a bullet even if style isn't found
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    ind.set(qn("w:left"),    "360")
    ind.set(qn("w:hanging"), "360")

    # Spacing
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"),  "40")


def _bullets(cell, items: List[str]):
    """Fill a cell with plain bullet points from a list of strings."""
    if not items:
        items = ["Not publicly available."]
    first_para = cell.paragraphs[0]
    for i, text in enumerate(items):
        if i == 0:
            para = first_para
        else:
            para = cell.add_paragraph()
        _bullet_para(para, text or "Not publicly available.")


def _bullets_mixed(cell, items: List[tuple]):
    """
    Fill a cell with a mix of bullet points and sub-labels.
    items: list of (text, is_sub_label) tuples.
    """
    if not items:
        items = [("Not publicly available.", False)]
    first_para = cell.paragraphs[0]
    for i, (text, is_sub_label) in enumerate(items):
        if i == 0:
            para = first_para
        else:
            para = cell.add_paragraph()
        _bullet_para(para, text or "", sub_label=is_sub_label)


# ══════════════════════════════════════════════
# Public entry point
# ══════════════════════════════════════════════

def build_profile_docx(data: dict, photo_bytes: Optional[bytes] = None) -> bytes:
    """
    Build a .docx profile from structured JSON data.

    Args:
        data:        Parsed JSON dict from profile_generator.research_prospect().
        photo_bytes: Raw image bytes (PNG / JPEG), or None.

    Returns:
        .docx file contents as bytes.
    """
    builder = DocxBuilder()
    profile_type = data.get("type", "individual")

    if profile_type == "company":
        return builder.build_company(data, photo_bytes)
    else:
        return builder.build_individual(data, photo_bytes)
